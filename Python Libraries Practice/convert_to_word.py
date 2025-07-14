from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.shared import qn

document = Document()
document.add_heading('Hello World', level=1)

p = document.add_paragraph('This is a paragraph in the Word document.')
p.add_run(' This text is bold.').bold = True
p.add_run(' This text is itlalic').italic = True
p.add_run(' This text is underlined.').underline = True
p.add_run(' This text is strikethrough.').strike = True

document.add_paragraph('This is one.', style='ListBullet')
document.add_paragraph('This is two.', style='ListBullet')
document.add_paragraph('This is three.', style='ListBullet')
document.add_paragraph('This is four.', style='ListBullet')

table_header = ["Name", "Age", "City", "Country"]

data = [
    ["Alice", 30, "New York", "USA"],
    ["Bob", 25, "Los Angeles", "USA"],
    ["Charlie", 35, "Chicago", "USA"],
    ["David", 28, "Houston", "USA"]
]

table = document.add_table(rows=1, cols=len(table_header)) # row for table heading

for i in range(len(table_header)): # adding data to table header
    table.rows[0].cells[i].text = table_header[i]
    
for name, age, city, country in data: # adding data to table rows
    row_cells = table.add_row().cells
    row_cells[0].text = name
    row_cells[1].text = str(age) # converting age to string
    row_cells[2].text = city
    row_cells[3].text = country
    
document.add_page_break()  # adding a page break

document.add_paragraph('This is a new page after the break.')

# function for adding a hyperlink
def add_hyperlink(paragraph, url, text=None): # there is no inbuilt function for adding hyperlinks in docx, so we create our own
    
    """
    Under the hood, python-docx is building and manipulating the raw WordprocessingML (that is, the XML schema that Word uses). The names you're seeing—like w:hyperlink, <w:r>, <w:rPr>, <w:u>, and the attribute r:id—are not Python syntax but XML element and attribute names defined by Microsoft's WordprocessingML spec.
    
    WordprocessingML uses XML namespaces, so every element really lives in a URI-qualified namespace. The prefix w: is shorthand for the WordprocessingML namespace URI.
    """
    if text is None:
        text = url
    # This function adds a hyperlink to a paragraph
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    # qn() simply expands the "w:…" shorthand into the full namespaced tag.
    
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id, )

    # Create a w:r element
    new_run = OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = OxmlElement('w:rPr')
    
    # adding underline to the hyperlink
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')  # setting underline style
    rPr.append(u)
    
    # adding color to the hyperlink
    color = OxmlElement('w:color')
    color.set(qn('w:val'), '0000FF')  # setting color to blue
    rPr.append(color)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


document.add_picture("Python Libraries Practice/thumb16.jpg", width=Document().sections[0].page_width / 2)  # adding an image with half the page width

github = document.add_paragraph(style='Normal')  # adding a normal paragraph for hyperlink
add_hyperlink(github, 'https://www.github.com', 'GitHub')  # adding github hyperlink


document.save('Python Libraries Practice/test.docx')